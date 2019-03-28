#!/usr/bin/python
#coding=utf-8
from jieba import lcut
from tqdm import tqdm
from imageio import imread
from docx import Document
from win32com.client import Dispatch,constants
from wordcloud import WordCloud
import os
from time import sleep
from os import path

# for i in ['豪尔赛','诺德','精进','沉淀']:
	# jieba.add_word(i)
# stopwords = []

mask_name = 'd:\\mask.jpg'

def txt_file2wc(file_name,mask_name='d:\\mask.jpg'):
	'''文本文件转词云图片'''
	f = open(file_name,'r')
	t = f.read()
	t.replace('\n','')
	ls = lcut(t)
	f.close()
	txt = ' '.join(ls)
	mk = imread(mask_name)
	w = WordCloud(font_path='simhei.ttf',\
	mask=mk,width=1000,height =700,\
	background_color='white').generate(txt)
	output_jpg = path.splitext(file_name)[0] + '.jpg' 
	w.to_file(output_jpg)



def alone_word2wc(f):
	'''对单个doc文件进行词云转化'''
	a = Document(f)
	ls = a.paragraphs
	txt = ''
	for i in ls:
		i.text.replace('\n',' ')
		txt += i.text
	txt = txt.replace('\n',' ')
	txt = lcut(txt)
	ls = ' '.join(txt)
	mk = imread(mask_name)
	w = WordCloud(font_path='simhei.ttf',\
	max_words=200,stopwords='王志学',mask=mk,width=800,height =600,\
	background_color='white')
	w.generate(ls)
	output_jpg = path.splitext(f)[0] + '.jpg'
	w.to_file(output_jpg)


def doc2docx(file_name):
	'''转换doc文件为docx格式'''
	w = Dispatch('Word.Application')
	w.Visible = 0
	w.DisplayAlerts = 0
	doc = w.Documents.Open(file_name)
	newpath = path.splitext(file_name)[0] + '.docx'
	doc.SaveAs(newpath, 12, False, "", True, "", False, False, False, False)
	doc.Close()
	w.Quit()
	os.remove(file_name)
	return newpath
    
 
def wordfold2wc(fold,mask_name='d:\\mask.jpg'):
	'''文件夹下所有Word文件转换成一个词云'''
	txt = ''
	output_jpg = ''
	list_2 = []
	list_1 = os.listdir(fold)
	for i in range(len(list_1)):
		if list_1[i][-3:0] == 'doc': #如果文件夹中有doc文件则转化之
			doc2docx(path.join(fold,file_1[i]))
	for i in list_1:
		if i[-4:] == 'docx':
			list_2.append(i)
	for i in list_2:
		i = path.join(fold,i)
		ls = Document(i).paragraphs
		# print(type(ls))
		for s in ls:
			txt += s.text
	txt.replace('\n',' ') #删除所有回车空行
	ls = lcut(txt)
	ls = ' '.join(ls)
	mk = imread(mask_name)
	w = WordCloud(font_path='simhei.ttf',mask=mk,width=800,height =600,\
	background_color='white')
	w.generate(ls)
	fold_name = fold.split('\\')[-1] + '.jpg'
	output_jpg = path.join(fold,fold_name)
	w.to_file(output_jpg)


def ppt2wc(file_name,mask_name='d:\\mask.jpg'):
	'''PPT文件转词云'''
	t = ''
	txt = ''
	ppt = Dispatch('PowerPoint.Application')
	ppt.Visible = 1
	pptSel = ppt.Presentations.Open(file_name) # win32com.client.gencache.EnsureDispatch('PowerPoint.Application')
	slide_count = pptSel.Slides.Count #get the ppt's pages
	for i in range(1,slide_count + 1):
		shape_count = pptSel.Slides(i).Shapes.Count
		for j in range(1,shape_count + 1):
			if pptSel.Slides(i).Shapes(j).HasTextFrame:
				s = pptSel.Slides(i).Shapes(j).TextFrame.TextRange.Text
				t += s
	t.replace('\n',' ')
	ls = lcut(t)
	txt = ' '.join(ls)
	ppt.Quit()
	mk = imread(mask_name)
	w = WordCloud(font_path='simhei.ttf',mask=mk,width=800,height =600,\
	background_color='white')
	w.generate(txt)
	output_jpg = path.splitext(file_name)[0] + '.jpg'
	w.to_file(output_jpg)
	
	

ls = []
f_fold_name = input('请输入需要转换的文件夹或者文件')
if path.isdir(f_fold_name): #如果输入的是文件夹
	select_way = input('输入数字1生成每个文件的单独词云，不输入则默认生产整个文件夹的词云')
	if select_way == '1':
		for i in os.listdir(f_fold_name):
			i2 = path.join(f_fold_name,i)
			if i2 not in ls:
				ls.append(i2)
		for i in range(len(ls)):
			if ls[i].split('\\')[-1][-4:] == '.doc':
				doc2docx(ls[i])
		ls2 = os.listdir(f_fold_name)
		for i in tqdm(range(len(ls2))):
			sleep(0.1)
			# print("\r词云已完成{:.2f}%{:.<30}".format(((i+1)/(len(ls2)-1)*100),"▊"*int(((i+1)/(len(ls2)-1))*33.3)),end='')
			if ls2[i][-5:] == '.docx':
				i2 = path.join(f_fold_name,ls2[i])
				alone_word2wc(i2)			
	else: #如果输入的不是数字"1",则对整个文件夹内的doc文件进行生成
		fold = f_fold_name
		wordfold2wc(fold)
elif path.isfile(f_fold_name): #如果输入的是文件
	print('您输入的是文件路径')
	if f_fold_name.split('\\')[-1][-3:] == 'doc': #如果输入的是doc文件
		f_fold_name = doc2docx(f_fold_name)
		print('您输入的是{}已经转换完成。'.format(f_fold_name))
		alone_word2wc(f_fold_name)
	elif f_fold_name.split('\\')[-1][-3:] == 'txt': #如果输入的是txt文件
		txt_file2wc(f_fold_name)
	elif f_fold_name.split('\\')[-1][-4:] == 'docx': #如果输入的是docx文件
		alone_word2wc(f_fold_name)
	elif f_fold_name.split('\\')[-1][-4:] == 'pptx': #如果输入的是ppt文件
		ppt2wc(f_fold_name,mask_name='d:\\mask.jpg')
	elif f_fold_name[-4:] == 'pptx': #如果输入的是字符串文本
		print('您输入的是ppt文件')
		ppt2wc(f_fold_name,mask_name='d:\\mask.jpg')
	else:
		print('输入错误')


	
		
	
